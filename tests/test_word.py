"""测试 WordParser Word 解析器"""
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from src.parsers.word_parser import WordParser
from pathlib import Path


def test_parse_file():
    """测试解析 Word 文件"""
    print("=== 测试解析 Word 文件 ===")

    parser = WordParser()
    test_file = Path(__file__).parent.parent / "data" / "test.docx"

    result = parser.parse(str(test_file))

    print(f"文件名: {result.filename}")
    print(f"总块数: {result.total_chunks}")
    print(f"表格数量: {result.table_count}")
    print(f"解析器: {result.metadata.get('parser')}")
    print(f"原始文本长度: {len(result.raw_text)}")
    print(f"原始文本预览: {result.raw_text[:200]}...")

    # 验证结果
    assert result.filename == "test.docx"
    assert result.total_pages == 1  # Word 不分页
    assert result.total_chunks >= 0
    assert result.raw_text is not None
    assert len(result.raw_text) > 0

    # 验证分块
    if result.chunks:
        print(f"\n第一个分块:")
        print(f"  索引: {result.chunks[0]['index']}")
        print(f"  长度: {result.chunks[0]['length']}")
        print(f"  类型: {result.chunks[0]['type']}")
        print(f"  内容预览: {result.chunks[0]['content'][:100]}...")

    # 验证表格
    if result.tables:
        print(f"\n检测到 {len(result.tables)} 个表格")
        for i, table in enumerate(result.tables[:3]):  # 显示前3个表格
            print(f"  表格 {i + 1}: {table['rows']}行 x {table['cols']}列")

    print("[OK] Word 文件解析测试通过")


def test_parse_from_bytes():
    """测试从字节流解析 Word"""
    print("\n=== 测试从字节流解析 Word ===")

    parser = WordParser()
    test_file = Path(__file__).parent.parent / "data" / "test.docx"

    with open(test_file, 'rb') as f:
        docx_bytes = f.read()

    result = parser.parse_from_bytes(docx_bytes, "test.docx")

    print(f"文件名: {result.filename}")
    print(f"总块数: {result.total_chunks}")
    print(f"表格数量: {result.table_count}")
    print(f"原始文本长度: {len(result.raw_text)}")

    # 验证结果
    assert result.filename == "test.docx"
    assert result.raw_text is not None
    assert len(result.raw_text) > 0

    print("[OK] 字节流解析测试通过")


def test_supported_extensions():
    """测试支持的文件扩展名"""
    print("\n=== 测试支持的文件扩展名 ===")

    parser = WordParser()
    extensions = parser.supported_extensions()

    print(f"支持的扩展名: {extensions}")

    # 验证支持的扩展名
    assert ".docx" in extensions

    print("[OK] 扩展名测试通过")


def test_table_extraction():
    """测试表格提取功能"""
    print("\n=== 测试表格提取功能 ===")

    parser = WordParser()
    test_file = Path(__file__).parent.parent / "data" / "test.docx"

    result = parser.parse(str(test_file))

    print(f"提取到的表格数量: {result.table_count}")

    if result.tables:
        for i, table in enumerate(result.tables):
            print(f"\n表格 {i + 1}:")
            print(f"  页码: {table.get('page')}")
            print(f"  行数: {table.get('rows')}")
            print(f"  列数: {table.get('cols')}")
            # 跳过包含 emoji 的内容
            md = table.get('markdown', '')[:150]
            try:
                print(f"  Markdown 预览:\n{md}...")
            except UnicodeEncodeError:
                print(f"  Markdown 预览: (包含特殊字符，跳过显示)")

    print("[OK] 表格提取测试通过")


def test_table_to_dataframe():
    """测试将表格转换为 DataFrame"""
    print("\n=== 测试表格转 DataFrame ===")

    parser = WordParser()
    test_file = Path(__file__).parent.parent / "data" / "test.docx"

    result = parser.parse(str(test_file))

    if result.tables:
        from src.parsers.table_parser import TableProcessor

        table_processor = TableProcessor()

        for i, table in enumerate(result.tables[:2]):  # 只处理前2个表格
            markdown = table.get('markdown', '')
            df = table_processor.markdown_to_dataframe(markdown)

            print(f"\n表格 {i + 1} 转换结果:")
            if df is not None:
                print(f"  DataFrame 形状: {df.shape}")
                print(f"  列名: {list(df.columns)}")
            else:
                print("  转换失败或表格为空")

    print("[OK] 表格转 DataFrame 测试通过")


def test_custom_chunker():
    """测试使用自定义分块器"""
    print("\n=== 测试使用自定义分块器 ===")

    from src.parsers.chunker import TextChunker, ChunkStrategy

    custom_chunker = TextChunker(
        strategy=ChunkStrategy.SENTENCE,
        max_chunk_size=300,
        min_chunk_size=50
    )

    parser = WordParser(chunker=custom_chunker)
    test_file = Path(__file__).parent.parent / "data" / "test.docx"

    result = parser.parse(str(test_file))

    print(f"使用自定义分块器（按句子）")
    print(f"总块数: {result.total_chunks}")

    if result.chunks:
        # 区分文本块和表格块
        text_chunks = [c for c in result.chunks if c['type'] == 'text']
        table_chunks = [c for c in result.chunks if c['type'] == 'table']

        print(f"  文本块数: {len(text_chunks)}")
        print(f"  表格块数: {len(table_chunks)}")

    assert result.total_chunks >= 0

    print("[OK] 自定义分块器测试通过")


def test_chunks_with_tables():
    """测试分块中是否包含表格"""
    print("\n=== 测试分块中是否包含表格 ===")

    parser = WordParser()
    test_file = Path(__file__).parent.parent / "data" / "test.docx"

    result = parser.parse(str(test_file))

    if result.chunks:
        table_chunks = [c for c in result.chunks if c['type'] == 'table']

        print(f"总分块数: {len(result.chunks)}")
        print(f"表格块数: {len(table_chunks)}")

        if table_chunks:
            for i, chunk in enumerate(table_chunks[:2]):
                print(f"\n表格块 {i + 1}:")
                print(f"  长度: {chunk['length']}")
                print(f"  页码: {chunk.get('page')}")

    print("[OK] 表格分块测试通过")


def test_paragraph_extraction():
    """测试段落提取"""
    print("\n=== 测试段落提取 ===")

    parser = WordParser()
    test_file = Path(__file__).parent.parent / "data" / "test.docx"

    result = parser.parse(str(test_file))

    # 检查是否有双换行符（段落分隔）
    paragraph_count = result.raw_text.count('\n\n') + 1
    print(f"估算的段落数: {paragraph_count}")
    print(f"原始文本长度: {len(result.raw_text)}")

    print("[OK] 段落提取测试通过")


def test_empty_document():
    """测试空文档处理"""
    print("\n=== 测试空文档处理 ===")

    from docx import Document
    from io import BytesIO

    # 创建一个空的 Word 文档
    doc = Document()
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    parser = WordParser()
    result = parser.parse_from_bytes(stream.read(), "empty.docx")

    print(f"空 Word 文档结果:")
    print(f"  文件名: {result.filename}")
    print(f"  文本长度: {len(result.raw_text)}")
    print(f"  块数: {result.total_chunks}")
    print(f"  表格数: {result.table_count}")

    print("[OK] 空文档处理测试通过")
