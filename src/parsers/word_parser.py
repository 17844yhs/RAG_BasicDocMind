from pathlib import Path
from typing import List, Dict, Optional

import docx   # python-docx

from src.parsers.base import BaseParser,ParseResult
from src.parsers.chunker import TextChunker
from src.parsers.cleaner import TextCleaner
from src.parsers.table_parser import TableProcessor


class WordParser(BaseParser):
    def __init__(self,
                 chunker:Optional[TextChunker] =None,
                 cleaner:Optional[TextCleaner] =None,
                 table_processor:Optional[TableProcessor] =None
                 ):

        # 初始化组件
        self.chunker = chunker if chunker else TextChunker()
        self.cleaner = cleaner if cleaner else TextCleaner()
        self.table_processor = table_processor if table_processor else TableProcessor()
    
    def parse(self, file_path:str) -> ParseResult:
        '''
        通过文件路径解析文件
        参数：
            file_path: 文件路径
        '''
        # 创建Docx对象
        doc  = docx.Document(file_path)

        # 提取数据
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        raw_text = '\n\n'.join(paragraphs)

        # 提取表格
        tables = self._extract_tables(doc)

        # 处理数据
        cleaned_text = self.cleaner.clean(raw_text)
        merged_tables = self.table_processor.merge_tables(tables)
        chunks = self.chunker.split(cleaned_text, merged_tables)

        return ParseResult(
            filename=Path(file_path).name,
            total_pages=1,  # 不分页
            total_chunks=len(chunks),
            table_count=len(merged_tables),
            chunks=chunks,
            raw_text=raw_text,
            tables=tables,
            metadata={'parser':'python-docx','filename':Path(file_path).name}
        )

    def parse_from_bytes(self, file_bytes, filename):
        '''
        通过字节流解析文件
        参数：
            file_bytes: 文件字节流
            filename: 文件名
        '''
        from io import BytesIO
        stream = BytesIO(file_bytes)
        doc = docx.Document(stream)

        # 提取数据
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        raw_text = '\n\n'.join(paragraphs)

        # 提取表格
        tables = self._extract_tables(doc)

        # 处理数据
        cleaned_text = self.cleaner.clean(raw_text)
        merged_tables = self.table_processor.merge_tables(tables)
        chunks = self.chunker.split(cleaned_text, merged_tables)

        return ParseResult(
            filename=filename,
            total_pages=1,  # 不分页
            total_chunks=len(chunks),
            table_count=len(merged_tables),
            chunks=chunks,
            raw_text=raw_text,
            tables=tables,
            metadata={'parser':'python-docx','filename':filename}
        )
    
    def supported_extensions(self) -> list[str]:
        '''
        支持的文件扩展名
        '''
        return ['.docx','.DOCX']
    
    def _extract_tables(self,doc) -> List[Dict]:
        '''
        提取表格,自己拼接成md格式
        '''
        tables = []
        for table_index,table in  enumerate(doc.tables):
            # 转换markdown
            md_rows = []
            for row_index,row in enumerate(table.rows):
                # 获取列
                cells = [cell.text.strip() for cell in row.cells]
                # | col1 | col2 | col3 |
                md_rows.append('| ' + ' | '.join(cells) + ' |')
                # 增加表头
                if row_index == 0:
                    md_rows.append( '| ' + ' | '.join(['---']*len(cells)) + ' |')
            tables.append({
                'page':1,
                'rows':len(table.rows),
                'cols':len(table.columns),
                'markdown':'\n'.join(md_rows)
            })
        return tables
